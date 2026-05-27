
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_document():
    doc = Document()

    # Estilo base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Título Principal
    title = doc.add_heading('Descripción Integral - Plataforma Turismo Chile MVP', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # INTRODUCCIÓN
    doc.add_heading('1. Introducción al Proyecto', level=1)
    doc.add_paragraph(
        'La Plataforma Turismo Chile es una aplicación de "Paquetes Dinámicos" diseñada para automatizar la búsqueda y combinación '
        'de vuelos y hoteles en los 10 destinos más importantes de Chile. El objetivo es proporcionar al usuario una visión clara '
        'del costo total de un viaje, utilizando datos en tiempo real siempre que sea posible.'
    )

    # ARQUITECTURA TÉCNICA
    doc.add_heading('2. Arquitectura y Componentes', level=1)
    
    doc.add_heading('Frontend (Next.js 16 + Tailwind)', level=2)
    doc.add_paragraph(
        'Es la capa de presentación. Utiliza componentes modernos de React y Framer Motion para una experiencia premium. '
        'Su lógica principal reside en:'
    )
    doc.add_paragraph('• SearchBar: Captura el destino y las fechas, validando que la fecha de ida sea anterior a la de vuelta.', style='List Bullet')
    doc.add_paragraph('• Página de Resultados (/buscar): Realiza una petición al API Proxy y muestra las tarjetas de paquetes con animaciones de carga.', style='List Bullet')
    doc.add_paragraph('• API Proxy: Actúa como intermediario para evitar problemas de CORS y centralizar las peticiones al backend de Python.', style='List Bullet')

    doc.add_heading('Backend (FastAPI + Python)', level=2)
    doc.add_paragraph(
        'Es el cerebro del sistema. Gestiona la orquestación de datos y el motor de búsqueda. Sus funciones clave son:'
    )
    doc.add_paragraph('• Endpoint /api/buscar: Recibe la solicitud, verifica la caché en la base de datos y, si no hay datos recientes, activa el scraper.', style='List Bullet')
    doc.add_paragraph('• Motor de Scraping: Ejecuta navegadores controlados por código para extraer precios de sitios externos.', style='List Bullet')
    doc.add_paragraph('• Lógica de Negocio: Combina los precios de vuelos y hoteles, aplica márgenes (si existen) y genera un paquete único.', style='List Bullet')

    doc.add_heading('Persistencia (PostgreSQL 16)', level=2)
    doc.add_paragraph(
        'Almacena toda la información crítica en 6 tablas normalizadas:'
    )
    doc.add_paragraph('• Destinos: Metadata de los 10 lugares (IATA, descripción, coordenadas).', style='List Bullet')
    doc.add_paragraph('• Vuelos/Hoteles Scrapeados: Almacena los resultados individuales para auditoría.', style='List Bullet')
    doc.add_paragraph('• Caché de Búsquedas: Evita hacer scraping repetitivo guardando resultados por 2 horas según destino y fecha.', style='List Bullet')

    # LÓGICA DE FUNCIONAMIENTO
    doc.add_heading('3. Lógica de Funcionamiento (El "Flujo")', level=1)
    doc.add_paragraph(
        'Cuando un usuario busca un viaje, el sistema sigue este orden lógico:'
    )
    
    step1 = doc.add_paragraph()
    step1.add_run('Paso 1: Validación y Caché. ').bold = True
    step1.add_run('El backend genera una firma única (SHA-256) para la búsqueda. Si esa búsqueda se realizó hace menos de 2 horas, entrega los resultados guardados instantáneamente.')

    step2 = doc.add_paragraph()
    step2.add_run('Paso 2: Activación del Scraper. ').bold = True
    step2.add_run('Si no hay caché, se inicia un hilo (thread) separado. Esto es vital para que el servidor no se bloquee mientras el navegador "fantasma" busca los precios.')

    step3 = doc.add_paragraph()
    step3.add_run('Paso 3: Extracción Real (Playwright). ').bold = True
    step3.add_run('El sistema navega a Booking.com y LATAM Airlines. Lee el código HTML de las páginas, busca los selectores de precio y extrae los valores numéricos.')

    step4 = doc.add_paragraph()
    step4.add_run('Paso 4: Inteligencia de Fallback. ').bold = True
    step4.add_run('Si un sitio bloquea el acceso (como suele ocurrir con LATAM), el sistema no falla. En su lugar, utiliza una base de datos de precios referenciales para mostrar un paquete estimado, asegurando que el usuario siempre tenga información útil.')

    # MOTOR DE SCRAPING
    doc.add_heading('4. El Motor de Scraping (Detalle Técnico)', level=1)
    doc.add_paragraph(
        'El motor fue rediseñado para ser "Resiliente y Honesto".'
    )
    doc.add_paragraph('• Estabilidad en Windows: Usamos ThreadPoolExecutor para ejecutar Playwright en modo síncrono. Esto evita errores de incompatibilidad con el bucle de eventos de Windows.', style='List Bullet')
    doc.add_paragraph('• Transparencia de Datos: El sistema marca cada resultado con una etiqueta (scraped o referencial). El frontend usa esta etiqueta para mostrar banners informativos al usuario.', style='List Bullet')
    doc.add_paragraph('• URLs Dinámicas: Genera enlaces de reserva directos que incluyen las fechas y el destino exacto para que el usuario pueda comprar con un clic.', style='List Bullet')

    # CAMBIOS RECIENTES
    doc.add_heading('5. Cambios Clave y Mejoras de Estabilidad', level=1)
    doc.add_paragraph('• Fix de Tipado DATE: Corrección en la comunicación con PostgreSQL para manejar fechas nativas de Python.', style='List Bullet')
    doc.add_paragraph('• Ruta de LATAM: Actualización masiva de selectores y estructura de URL para cumplir con los estándares de 2026.', style='List Bullet')
    doc.add_paragraph('• Automatización run_all.ps1: Script unificado que reduce el tiempo de arranque del entorno de 5 minutos a 10 segundos.', style='List Bullet')

    # GUÍA DE OPERACIÓN
    doc.add_heading('6. Guía de Operación', level=1)
    doc.add_paragraph('Para ejecutar la plataforma completa:')
    doc.add_paragraph('1. Abrir PowerShell como Administrador.', style='List Number')
    doc.add_paragraph('2. Navegar a la carpeta del proyecto.', style='List Number')
    doc.add_paragraph('3. Ejecutar: .\\run_all.ps1', style='List Number')
    doc.add_paragraph('Esto abrirá las ventanas correspondientes para el Frontend (puerto 3000) y el Backend (puerto 8001).', style='List Number')

    # Finalizar
    save_path = os.path.join(os.path.expanduser('~'), 'Downloads', 'Descripcion_Total_Turismo_Chile_MVP.docx')
    doc.save(save_path)
    print(f'Documento guardado en: {save_path}')

if __name__ == "__main__":
    create_document()
