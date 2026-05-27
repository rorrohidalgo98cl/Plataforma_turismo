import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def parse_destinos(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Extract the DESTINOS array content
    match = re.search(r'export const DESTINOS: Destino\[\] = \[(.*?)\]\s*export function', content, re.DOTALL)
    if not match:
        # Fallback regex if footer differs
        match = re.search(r'export const DESTINOS: Destino\[\] = \[(.*?)\];\s*$', content, re.DOTALL)
        if not match:
            match = re.search(r'export const DESTINOS: Destino\[\] = \[(.*)', content, re.DOTALL)
            
    items_text = match.group(1) if match else content

    # Find all blocks between { and }
    blocks = re.findall(r'\{([^\}]+)\}', items_text)
    
    destinos = []
    for block in blocks:
        destino = {}
        
        # Helper to extract string fields
        def get_str(field):
            m = re.search(rf'{field}:\s*[\'"]([^\'"]+)[\'"]', block)
            return m.group(1) if m else ''

        # Helper to extract bool
        def get_bool(field):
            m = re.search(rf'{field}:\s*(true|false)', block)
            return m.group(1) == 'true' if m else False

        # Helper to extract number
        def get_num(field):
            m = re.search(rf'{field}:\s*(\d+)', block)
            return int(m.group(1)) if m else 0

        # Helper to extract tags
        def get_tags():
            m = re.search(r'tags:\s*\[(.*?)\]', block)
            if m:
                tags_str = m.group(1)
                return [t.strip().strip('\'"') for t in tags_str.split(',') if t.strip()]
            return []

        slug = get_str('slug')
        if not slug:
            continue
            
        destino['slug'] = slug
        destino['nombre'] = get_str('nombre')
        destino['region'] = get_str('region')
        destino['tagline'] = get_str('tagline')
        destino['descripcion'] = get_str('descripcion')
        destino['codigoIATA'] = get_str('codigoIATA')
        destino['precioDesde'] = get_num('precioDesde')
        destino['destacado'] = get_bool('destacado')
        destino['tags'] = get_tags()
        
        destinos.append(destino)
        
    return destinos

def generar_documento_word(destinos, output_path):
    doc = Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(51, 51, 51)

    # Title
    title = doc.add_heading('CATÁLOGO E INVENTARIO GENERAL DE DESTINOS', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_font = title.runs[0].font
    title_font.name = 'Calibri'
    title_font.size = Pt(24)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 102, 204)
    
    subtitle = doc.add_paragraph('Plataforma de Turismo Chile - MVP & Paquetes Dinámicos')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.italic = True
    subtitle.runs[0].font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph() # Spacing

    # Intro
    intro = doc.add_paragraph(
        f'El presente documento detalla el inventario completo de los {len(destinos)} destinos y atractivos turísticos '
        'actualmente implementados en la base de datos de la Plataforma de Turismo Chile. Este catálogo abarca tanto los '
        'destinos principales (cabeceras de paquetes dinámicos) como los atractivos regionales específicos que enriquecen '
        'la oferta turística del sistema.'
    )
    
    # Separar Destacados (Principales) y Atractivos
    principales = [d for d in destinos if d['destacado']]
    atractivos = [d for d in destinos if not d['destacado']]
    
    # SECCIÓN 1: DESTINOS PRINCIPALES
    h1 = doc.add_heading(f'1. Destinos Principales ({len(principales)} Destinos)', level=1)
    h1.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph('Los destinos principales actúan como los centros de operación y búsqueda para la conformación de paquetes dinámicos (vuelo + hotel).')
    
    for d in principales:
        h2 = doc.add_heading(f'• {d["nombre"]} ({d["region"]})', level=2)
        h2.runs[0].font.color.rgb = RGBColor(0, 153, 102)
        
        p_tag = doc.add_paragraph()
        p_tag.add_run('Slogan: ').bold = True
        p_tag.add_run(f'"{d["tagline"]}"').italic = True
        
        p_desc = doc.add_paragraph()
        p_desc.add_run('Descripción: ').bold = True
        p_desc.add_run(d['descripcion'])
        
        p_info = doc.add_paragraph()
        p_info.add_run('Aeropuerto (IATA): ').bold = True
        p_info.add_run(f'{d["codigoIATA"]} | ')
        p_info.add_run('Precio Referencial Desde: ').bold = True
        precio_fmt = f'${d["precioDesde"]:,}'.replace(',', '.')
        p_info.add_run(f'{precio_fmt} CLP | ')
        p_info.add_run('Etiquetas: ').bold = True
        p_info.add_run(', '.join(d['tags']).title())
        
        doc.add_paragraph() # Spacing

    # SECCIÓN 2: ATRACTIVOS REGIONALES (TABLA RESUMEN)
    h1_attr = doc.add_heading(f'2. Atractivos y Sub-destinos Regionales ({len(atractivos)} Atractivos)', level=1)
    h1_attr.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph('A continuación se listan los atractivos turísticos específicos clasificados por región, los cuales complementan la oferta de los destinos principales en las búsquedas y recomendaciones:')
    
    # Agrupar atractivos por región
    regiones = {}
    for a in atractivos:
        r = a['region']
        if r not in regiones:
            regiones[r] = []
        regiones[r].append(a)
        
    for region, lista in sorted(regiones.items()):
        h2_reg = doc.add_heading(f'Región: {region} ({len(lista)} atractivos)', level=2)
        h2_reg.runs[0].font.color.rgb = RGBColor(102, 102, 102)
        
        # Tabla
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        # Encabezados
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Atractivo / Nombre'
        hdr_cells[1].text = 'IATA'
        hdr_cells[2].text = 'Descripción / Slogan'
        hdr_cells[3].text = 'Etiquetas'
        
        # Formato de encabezado
        widths = [Inches(2.0), Inches(0.8), Inches(2.7), Inches(1.0)]
        for i, cell in enumerate(hdr_cells):
            cell.width = widths[i]
            set_cell_background(cell, '0066CC')
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)
                    r.font.size = Pt(10)
                    
        # Filas
        for row_idx, a in enumerate(lista):
            row_cells = table.add_row().cells
            row_cells[0].text = a['nombre']
            row_cells[1].text = a['codigoIATA']
            row_cells[2].text = a['tagline']
            row_cells[3].text = ', '.join(a['tags']).title()
            
            # Formato celdas y cebrado
            bg_color = 'F9F9F9' if row_idx % 2 == 1 else 'FFFFFF'
            for i, cell in enumerate(row_cells):
                cell.width = widths[i]
                set_cell_background(cell, bg_color)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9.5)
                        
        doc.add_paragraph() # Spacing

    doc.save(output_path)
    print(f'Documento Word generado exitosamente en: {output_path}')

def generar_markdown_artifact(destinos, output_path):
    principales = [d for d in destinos if d['destacado']]
    atractivos = [d for d in destinos if not d['destacado']]
    
    md = [
        "# 🗺️ Catálogo e Inventario General de Destinos",
        "**Plataforma de Turismo Chile - MVP & Paquetes Dinámicos**\n",
        f"> [!IMPORTANT]",
        f"> **Resumen Ejecutivo:** La plataforma cuenta actualmente con **{len(destinos)} destinos totales** implementados en su base de datos (`destinos.ts`), divididos en **{len(principales)} Destinos Principales** (cabeceras de paquetes dinámicos) y **{len(atractivos)} Atractivos Regionales** (sub-destinos para exploración y armado de itinerarios).\n",
        "---",
        "\n## 🌟 1. Destinos Principales del MVP",
        "Los destinos principales actúan como los centros de operación y búsqueda para la conformación de paquetes dinámicos (vuelo + hotel).\n"
    ]
    
    for d in principales:
        precio_fmt = f'${d["precioDesde"]:,}'.replace(',', '.')
        tags_badge = ' '.join([f'`#{t}`' for t in d['tags']])
        md.append(f"### 📍 {d['nombre']} ({d['region']})")
        md.append(f"**Slogan:** *\"{d['tagline']}\"*")
        md.append(f"- **Descripción:** {d['descripcion']}")
        md.append(f"- **Aeropuerto más cercano (IATA):** ✈️ `{d['codigoIATA']}`")
        md.append(f"- **Precio Referencial Paquete:** 💰 **{precio_fmt} CLP**")
        md.append(f"- **Etiquetas:** {tags_badge}\n")
        
    md.append("---")
    md.append("\n## 🏞️ 2. Atractivos y Sub-destinos Regionales")
    md.append("A continuación se listan los atractivos turísticos específicos clasificados por región, los cuales complementan la oferta de los destinos principales en las búsquedas y recomendaciones:\n")
    
    regiones = {}
    for a in atractivos:
        r = a['region']
        if r not in regiones:
            regiones[r] = []
        regiones[r].append(a)
        
    for region, lista in sorted(regiones.items()):
        md.append(f"### 📌 Región: {region} ({len(lista)} atractivos)")
        md.append("| Atractivo / Nombre | IATA | Slogan / Descripción | Etiquetas |")
        md.append("| :--- | :---: | :--- | :--- |")
        for a in lista:
            tags_str = ', '.join(a['tags']).title()
            md.append(f"| **{a['nombre']}** | `{a['codigoIATA']}` | {a['tagline']} | `{tags_str}` |")
        md.append("\n")
        
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(md))
    print(f'Artefacto Markdown generado exitosamente en: {output_path}')

if __name__ == '__main__':
    base_dir = os.path.dirname(os.path.abspath(__file__))
    destinos_ts_path = os.path.join(base_dir, 'frontend', 'src', 'data', 'destinos.ts')
    
    destinos = parse_destinos(destinos_ts_path)
    print(f'Se extrajeron {len(destinos)} destinos desde {destinos_ts_path}')
    
    # Rutas de salida
    word_out = os.path.join(os.path.expanduser('~'), 'Downloads', 'Inventario_Destinos_Turismo_Chile.docx')
    md_out = os.path.join(base_dir, 'inventario_destinos.md')
    
    generar_documento_word(destinos, word_out)
    generar_markdown_artifact(destinos, md_out)
