from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        p.add_run(bold_prefix).bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def create_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('Plataforma Turismo Chile MVP\nDocumentación Técnica y Arquitectura', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Introducción
    add_heading(doc, '1. Introducción al Proyecto', level=1)
    doc.add_paragraph(
        'La Plataforma Turismo Chile es una aplicación moderna de "Paquetes Dinámicos" ('
        'Dynamic Packaging) enfocada en el mercado nacional. El sistema combina vuelos y '
        'hoteles para 10 de los principales destinos turísticos de Chile, ofreciendo a los '
        'usuarios una experiencia unificada con precios en tiempo real o referenciales '
        'altamente transparentes.'
    )

    # 2. Arquitectura General
    add_heading(doc, '2. Arquitectura del Sistema (3 Capas)', level=1)
    doc.add_paragraph(
        'El proyecto está diseñado bajo una arquitectura de microservicios dividida en '
        'tres capas principales, lo que garantiza escalabilidad y separación de '
        'responsabilidades:'
    )
    
    add_heading(doc, '2.1. Frontend: Next.js 16 + Turbopack', level=2)
    doc.add_paragraph(
        'Actúa como la capa de presentación, ofreciendo una experiencia de usuario '
        'rápida, interactiva y responsiva.'
    )
    add_bullet(doc, 'Tecnologías:', ' React 19, TailwindCSS, Framer Motion.')
    add_bullet(doc, 'Lógica Core:', ' Implementa una SearchBar compleja que maneja validaciones de fechas. La página de resultados (/buscar) se comunica con un API Proxy local en Next.js, evitando exponer el backend de Python directamente al cliente y solucionando problemas de CORS.')
    
    add_heading(doc, '2.2. Backend Orquestador: FastAPI (Python 3.11)', level=2)
    doc.add_paragraph(
        'Es el cerebro de la operación. Procesa las solicitudes, gestiona la caché, '
        'y orquesta la recolección de datos.'
    )
    add_bullet(doc, 'Motor de Scraping (Playwright):', ' Corre en un hilo separado (ThreadPool) para no bloquear el servidor principal.')
    add_bullet(doc, 'Fallback Inteligente:', ' Si el scraping falla, calcula un precio estimado referencial basado en promedios históricos para no dejar al usuario sin opciones.')

    add_heading(doc, '2.3. Persistencia: PostgreSQL 16.8', level=2)
    doc.add_paragraph(
        'Base de datos relacional robusta que mantiene el estado del sistema a través de 6 tablas normalizadas:'
    )
    add_bullet(doc, 'Tablas estáticas:', ' "destinos" (contiene los 10 destinos predefinidos).')
    add_bullet(doc, 'Tablas dinámicas:', ' "vuelos_scrapeados", "hoteles_scrapeados", y "paquetes_armados" que guardan los datos recolectados.')
    add_bullet(doc, 'Caché:', ' "cache_busquedas" implementa un sistema de TTL de 2 horas basado en un hash SHA-256 de los parámetros de búsqueda, reduciendo la carga del scraper dramáticamente.')

    # 3. Diagnóstico y Scraping
    add_heading(doc, '3. Motor de Scraping: Diagnóstico y Evolución', level=1)
    doc.add_paragraph(
        'El componente más complejo del sistema es la obtención de datos en vivo. Recientemente '
        'sufrió una reestructuración completa debido a problemas técnicos graves en entornos Windows.'
    )
    
    add_heading(doc, '3.1. El Problema Crítico (Event Loop)', level=2)
    doc.add_paragraph(
        'El scraping asíncrono (async_playwright) nunca funcionó originalmente. Esto se debía a que '
        'Uvicorn (el servidor de FastAPI) en modo recarga (reload) utiliza el ProactorEventLoop de Windows, '
        'el cual presentaba un error fatal de NotImplementedError al intentar levantar subprocesos.'
    )
    
    add_heading(doc, '3.2. La Solución Arquitectónica', level=2)
    doc.add_paragraph(
        'La solución consistió en migrar el motor de scraping asíncrono a uno síncrono (sync_playwright) '
        'y ejecutarlo dentro de un ThreadPoolExecutor. Esto aísla el proceso del event loop de asyncio, '
        'evitando completamente el bug de Windows.'
    )

    add_heading(doc, '3.3. Estado Actual de Fuentes de Datos', level=2)
    add_bullet(doc, 'Playwright:', ' ✅ Ahora se inicializa, lanza Chromium y navega sin interrupciones.')
    add_bullet(doc, 'Booking.com:', ' ✅ Scraping exitoso. El sistema captura precios reales de hoteles e inyecta los nombres verídicos en el paquete.')
    add_bullet(doc, 'LATAM Airlines:', ' ⚠️ Alta protección Anti-Bot. LATAM bloquea los navegadores automatizados. Para manejar esto éticamente, el sistema fue reescrito para evitar mostrar números aleatorios (falsos). En su lugar, muestra un "precio referencial transparente" e invita al usuario a verificar el valor real directamente en la aerolínea.')
    add_bullet(doc, 'URLs de Reserva:', ' ✅ Completamente funcionales. Dirigen al usuario al funnel de compra exacto.')

    # 4. Cambios Importantes Recientes
    add_heading(doc, '4. Resumen de Cambios Estructurales Recientes', level=1)
    add_bullet(doc, 'Fix Tipado DATE (PostgreSQL):', ' Se corrigió un error donde se pasaban strings (ej: "2026-05-10") a la base de datos. Ahora se instancian como objetos date() nativos de Python, previniendo crashes en el driver asyncpg.')
    add_bullet(doc, 'Flexibilidad de Parámetros de Búsqueda:', ' Se agregaron aliases lógicos en el frontend. Ahora es indiferente si la ruta recibe "?ida=" o "?fecha_ida=", garantizando que la navegación entre componentes no rompa la búsqueda.')
    add_bullet(doc, 'Mantenimiento de Endpoints Externos:', ' Se corrigió la URL objetivo de LATAM (de "oferta-vuelos" a "ofertas-vuelos") y se adaptó la inyección de fechas al estándar ISO con milisegundos que ahora exige la aerolínea.')
    add_bullet(doc, 'Orquestación de Entorno (run_all.ps1):', ' Se creó un script en PowerShell que automatiza el levantamiento de los 3 microservicios (PostgreSQL, FastAPI y Next.js) con un solo comando, optimizando el ciclo de desarrollo.')

    # Save
    save_path = os.path.join(os.path.expanduser('~'), 'Downloads', 'Documentacion_Completa_Plataforma_Turismo.docx')
    doc.save(save_path)
    print(f'Documento guardado exitosamente en: {save_path}')

if __name__ == "__main__":
    create_document()
